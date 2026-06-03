# -*- coding: utf-8 -*-
"""
sdl_dwr_step2_generator.py

Generate Step 2 SDL Daily Work Report DOCX from the daily Excel file.
Designed for a small Flask app on Namecheap/cPanel shared hosting.

Required packages:
    openpyxl
    python-docx

Expected server files in the same folder as this script/app:
    Daily Work Report Step 2 Template - enTop v1.0.0.docx

Public function:
    generate_dwr_step2(excel_path, template_path=None, output_dir=None) -> str
"""

from __future__ import annotations

import html
import io
import os
import re
import shutil
import struct
import tempfile
import zipfile
from collections import defaultdict
from datetime import date, datetime
from typing import Any, Dict, Iterable, List, Tuple

from docx import Document
from docx.table import _Cell
from openpyxl import load_workbook

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_TEMPLATE = os.path.join(BASE_DIR, "Daily Work Report Step 2 Template - enTop v1.0.0.docx")
DEFAULT_OUTPUT_DIR = os.path.join(BASE_DIR, "generated_reports")

BRACKETS = [
    (0, 500),
    (501, 2000),
    (2001, 5000),
    (5001, 10000),
    (10001, 15000),
    (15001, 25000),
    (25001, 50000),
    (50001, 75000),
    (75001, float("inf")),
]


def _safe_float(value: Any, default: float = 0.0) -> float:
    if value is None or value == "":
        return default
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _safe_int(value: Any, default: int = 0) -> int:
    if value is None or value == "":
        return default
    try:
        return int(round(float(value)))
    except (TypeError, ValueError):
        return default


def _normalize_job_no(raw: Any) -> str:
    if raw is None:
        return ""
    if isinstance(raw, (int, float)):
        return str(int(raw))
    return str(raw).strip()


def _normalize_customer(raw: Any) -> str:
    return " ".join(str(raw).split()) if raw else ""


def _normalize_status(raw: Any) -> str:
    return str(raw).strip().upper() if raw is not None else ""


def _normalize_grp(raw: Any) -> str:
    return str(raw).strip() if raw is not None else ""


def _fmt_currency(v: Any) -> str:
    return f"${_safe_float(v):,.2f}"


def _fmt_pct_decimal(v: float) -> str:
    return f"{v * 100:.2f}%"


def _excel_serial_or_date(value: Any) -> date:
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    try:
        return datetime.strptime(str(value)[:10], "%Y-%m-%d").date()
    except Exception:
        return date.today()


def _read_excel_data(excel_path: str) -> Dict[str, Any]:
    wb = load_workbook(excel_path, data_only=True)
    if "Summary" not in wb.sheetnames or "Production_Center" not in wb.sheetnames:
        raise ValueError("Input workbook must contain 'Summary' and 'Production_Center' sheets.")

    ws_sum = wb["Summary"]
    ws_pc = wb["Production_Center"]
    dwr_no = str(ws_sum["B1"].value).strip()
    report_date = _excel_serial_or_date(ws_sum["C1"].value)

    jobs: List[Dict[str, Any]] = []
    for r in range(8, min(ws_pc.max_row, 108) + 1):
        job_type = str(ws_pc.cell(r, 4).value or "").strip().upper()
        job_no = _normalize_job_no(ws_pc.cell(r, 5).value)
        if not job_type and not job_no:
            continue
        value = _safe_float(ws_pc.cell(r, 11).value)
        customer = _normalize_customer(ws_pc.cell(r, 12).value)
        grp_work = _normalize_grp(ws_pc.cell(r, 15).value)
        rm_time = _safe_float(ws_pc.cell(r, 21).value)
        status = _normalize_status(ws_pc.cell(r, 25).value)
        if grp_work.lower() == "continued":
            continue
        jobs.append({
            "row": r,
            "type": job_type,
            "job_no": job_no,
            "value": value,
            "customer": customer,
            "grp_work": grp_work,
            "rm_time": rm_time,
            "status": status,
        })

    active = [j for j in jobs if j["status"] in {"C", "IP"}]
    active_o = [j for j in active if j["type"] == "O"]
    active_q = [j for j in active if j["type"] == "Q"]

    ord_val = sum(j["value"] for j in active_o)
    q_val = sum(j["value"] for j in active_q)
    total_val = ord_val + q_val
    ord_hrs = sum(j["rm_time"] for j in active_o)
    q_hrs = sum(j["rm_time"] for j in active_q)
    all_hrs = ord_hrs + q_hrs

    return {
        "dwr_no": dwr_no,
        "date": report_date,
        "jobs": jobs,
        "active": active,
        "active_o": active_o,
        "active_q": active_q,
        "ord_val": ord_val,
        "q_val": q_val,
        "total_val": total_val,
        "ord_hrs": ord_hrs,
        "q_hrs": q_hrs,
        "all_hrs": all_hrs,
        "active_o_count": len(active_o),
        "active_q_count": len(active_q),
        "resource_rows": [[ws_sum.cell(r, c).value for c in range(3, 10)] for r in range(3, 7)],
    }


def _make_leaderboard(jobs: List[Dict[str, Any]], sort_key: str, top_n: int = 3) -> List[Tuple[str, float, int, float]]:
    cust = defaultdict(lambda: {"value": 0.0, "count": 0, "time": 0.0})
    for j in jobs:
        customer = j["customer"] or "UNKNOWN CUSTOMER"
        cust[customer]["value"] += j["value"]
        cust[customer]["count"] += 1
        cust[customer]["time"] += j["rm_time"]
    ranked = sorted(cust.items(), key=lambda x: (-x[1][sort_key], x[0]))
    top = ranked[:top_n]
    rest = ranked[top_n:]
    rows = [(c, d["value"], int(d["count"]), d["time"]) for c, d in top]
    # Always return 4 rows. This keeps chart/table shape stable.
    while len(rows) < top_n:
        rows.append(("", 0.0, 0, 0.0))
    rows.append((
        "OTHER CUSTOMERS",
        sum(d["value"] for _, d in rest),
        int(sum(d["count"] for _, d in rest)),
        sum(d["time"] for _, d in rest),
    ))
    return rows


def _frequency(values: Iterable[float]) -> List[int]:
    result = []
    vals = list(values)
    for low, high in BRACKETS:
        if low == 0:
            result.append(sum(1 for v in vals if 0 <= v <= high))
        elif high == float("inf"):
            result.append(sum(1 for v in vals if v >= low))
        else:
            result.append(sum(1 for v in vals if low <= v <= high))
    return result


def _compute_metrics(data: Dict[str, Any]) -> Dict[str, Any]:
    active_o = data["active_o"]
    active_q = data["active_q"]
    metrics = dict(data)
    metrics.update({
        "ord_lb_count": _make_leaderboard(active_o, "count"),
        "ord_lb_value": _make_leaderboard(active_o, "value"),
        "ord_lb_time": _make_leaderboard(active_o, "time"),
        "q_lb_value": _make_leaderboard(active_q, "value"),
        "q_lb_time": _make_leaderboard(active_q, "time"),
        "ord_freq": _frequency(j["value"] for j in active_o),
        "q_freq": _frequency(j["value"] for j in active_q),
    })
    return metrics


def _set_cell_text(cell: _Cell, text: Any) -> None:
    text = "" if text is None else str(text)
    if not cell.paragraphs:
        cell.text = text
        return
    p = cell.paragraphs[0]
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for run in p.runs[1:]:
        run.text = ""
    for extra_p in cell.paragraphs[1:]:
        for run in extra_p.runs:
            run.text = ""


def _replace_in_paragraphs(paragraphs: Iterable[Any], replacements: Dict[str, str]) -> None:
    for p in paragraphs:
        full = "".join(run.text for run in p.runs)
        new_full = full
        for old, new in replacements.items():
            new_full = new_full.replace(old, new)
        if new_full != full and p.runs:
            p.runs[0].text = new_full
            for r in p.runs[1:]:
                r.text = ""


def _replace_everywhere(doc: Document, replacements: Dict[str, str]) -> None:
    _replace_in_paragraphs(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs, replacements)
    for section in doc.sections:
        for story in (section.header, section.footer):
            _replace_in_paragraphs(story.paragraphs, replacements)
            for table in story.tables:
                for row in table.rows:
                    for cell in row.cells:
                        _replace_in_paragraphs(cell.paragraphs, replacements)


def _patch_analysis_tables(docx_path: str, metrics: Dict[str, Any]) -> None:
    doc = Document(docx_path)
    report_date = metrics["date"]
    date_long = report_date.strftime("%d %B %Y").lstrip("0")
    date_upper = date_long.upper()
    date_short = report_date.strftime("%d/%m/%y")
    dwr = metrics["dwr_no"]

    replacements = {
        "DWR# 2026-90": f"DWR# {dwr}",
        "21/05/26": date_short,
        "21 May 2026": date_long,
        "[PUBLISH DATE]": date_upper,
        "[Publish Date]": date_upper,
        "……": date_long,
        ".....": date_long,
        "…… reative": f"{date_long} relative",
        "reative": "relative",
    }
    _replace_everywhere(doc, replacements)

    # Patch the two visible analysis tables by their header text.
    # Do not rely on fixed table indexes: the Step 2 template has only two visible
    # Word tables, while charts/embedded workbooks are separate OOXML parts.
    ord_val = metrics["ord_val"] or 1.0
    q_val = metrics["q_val"] or 1.0
    ord_lb_value = metrics["ord_lb_value"]
    q_lb_value = metrics["q_lb_value"]

    def find_table_by_header(header_text: str):
        header_text_norm = header_text.strip().lower()
        for table in doc.tables:
            if not table.rows:
                continue
            first_row_text = " | ".join(cell.text for cell in table.rows[0].cells).lower()
            if header_text_norm in first_row_text:
                return table
        return None

    def patch_customer_table(table, rows, total_label: str, pct_denominator: float, uppercase_top: bool = False) -> None:
        if table is None:
            raise ValueError(f"Could not find Step 2 analysis table with footer label: {total_label}")
        if len(table.rows) < 6 or len(table.columns) < 4:
            raise ValueError(f"Step 2 analysis table has unexpected shape: {len(table.rows)} rows x {len(table.columns)} columns")

        top_total_v = sum(r[1] for r in rows[:3])
        top_total_c = sum(r[2] for r in rows[:3])
        other = rows[3]
        output_rows = rows[:3] + [(total_label, top_total_v, top_total_c, 0.0), other]

        for idx, (name, value, count, _time) in enumerate(output_rows, start=1):
            display_name = name.upper() if uppercase_top and idx <= 3 and name else name
            pct = (value / pct_denominator) if pct_denominator else 0.0
            _set_cell_text(table.cell(idx, 0), display_name)
            _set_cell_text(table.cell(idx, 1), _fmt_currency(value))
            _set_cell_text(table.cell(idx, 2), str(int(count)))
            _set_cell_text(table.cell(idx, 3), _fmt_pct_decimal(pct))

    order_table = find_table_by_header("Ordering Customers")
    quote_table = find_table_by_header("Quoted Customers")

    # Ordering table is about Order customers only, so % of Value uses total active Order value.
    # Quoted table is about Quote customers only, so % of Value uses total active Quote value.
    patch_customer_table(order_table, ord_lb_value, "Top Ordering Customers Total", ord_val, uppercase_top=True)
    patch_customer_table(quote_table, q_lb_value, "Top Quoted Customers Total", q_val, uppercase_top=False)

    doc.save(docx_path)


def _xml_escape(v: Any) -> str:
    return html.escape(str(v), quote=False)


def _num(v: Any) -> str:
    f = _safe_float(v)
    # compact, stable numeric XML text
    return f"{f:.15g}"


def _replace_numref_values(xml: str, formula: str, new_vals: List[Any], force_format_code: str | None = None) -> str:
    def replace_block(m: re.Match) -> str:
        block = m.group(0)
        if f"<c:f>{formula}</c:f>" not in block:
            return block
        if force_format_code is not None:
            fmt = f"<c:formatCode>{force_format_code}</c:formatCode>"
        else:
            fc = re.search(r"<c:formatCode>[^<]*</c:formatCode>", block)
            fmt = fc.group(0) if fc else ""
        pts_xml = f'<c:ptCount val="{len(new_vals)}"/>' + "".join(
            f'<c:pt idx="{i}"><c:v>{_num(v)}</c:v></c:pt>' for i, v in enumerate(new_vals)
        )
        new_cache = f"<c:numCache>{fmt}{pts_xml}</c:numCache>"
        return re.sub(r"<c:numCache>.*?</c:numCache>", new_cache, block, flags=re.DOTALL)
    return re.sub(r"<c:numRef>.*?</c:numRef>", replace_block, xml, flags=re.DOTALL)


def _replace_strcache(xml: str, formula: str, new_names: List[str]) -> str:
    def repl(m: re.Match) -> str:
        block = m.group(0)
        if f"<c:f>{formula}</c:f>" not in block:
            return block
        pts_xml = f'<c:ptCount val="{len(new_names)}"/>' + "".join(
            f'<c:pt idx="{i}"><c:v>{_xml_escape(v)}</c:v></c:pt>' for i, v in enumerate(new_names)
        )
        return re.sub(r"<c:strCache>.*?</c:strCache>", f"<c:strCache>{pts_xml}</c:strCache>", block, flags=re.DOTALL)
    return re.sub(r"<c:strRef>.*?</c:strRef>", repl, xml, flags=re.DOTALL)


def _round_up_to_nearest(x: float, nearest: int = 50000) -> int:
    if x <= 0:
        return nearest
    return int(((x + nearest - 1) // nearest) * nearest)


def _patch_charts(docx_path: str, metrics: Dict[str, Any]) -> None:
    tmp = docx_path + ".charts.tmp"
    ord_val = metrics["ord_val"]
    q_val = metrics["q_val"]
    total_val = metrics["total_val"] or 1.0
    ord_hrs = metrics["ord_hrs"]
    q_hrs = metrics["q_hrs"]
    all_hrs = metrics["all_hrs"] or 1.0

    chart_updates: Dict[str, str] = {}
    with zipfile.ZipFile(docx_path, "r") as zin:
        for item in zin.infolist():
            if not re.match(r"word/charts/chart[1-7]\.xml$", item.filename):
                continue
            xml = zin.read(item.filename).decode("utf-8")
            if item.filename.endswith("chart1.xml"):
                xml = _replace_numref_values(xml, "Sheet1!$J$3:$J$5", [ord_val, q_val])
                xml = _replace_numref_values(xml, "Sheet1!$L$3:$L$5", [ord_val / total_val, q_val / total_val], "0.00%")
                xml = _replace_numref_values(xml, "Sheet1!$M$3:$M$5", [ord_hrs, q_hrs])
                xml = _replace_numref_values(xml, "Sheet1!$N$3:$N$5", [ord_hrs / all_hrs, q_hrs / all_hrs], "0.00%")
                xml = _replace_numref_values(xml, "Sheet1!$K$3:$K$5", [metrics["active_o_count"], metrics["active_q_count"]])
                y_max = _round_up_to_nearest(max(ord_val, q_val), 50000)
                xml = re.sub(r'<c:max val="[^"]+"/>', f'<c:max val="{y_max}"/>', xml, count=1)
                xml = re.sub(r'<c:dLblPos val="[^"]+"/>', '<c:dLblPos val="inEnd"/>', xml)
            elif item.filename.endswith("chart2.xml"):
                xml = _replace_numref_values(xml, "Sheet1!$Q$2:$Q$10", metrics["ord_freq"])
                xml = _replace_numref_values(xml, "Sheet1!$R$2:$R$10", metrics["q_freq"])
            elif item.filename.endswith("chart3.xml"):
                lb = metrics["ord_lb_count"]
                xml = _replace_strcache(xml, "Summary_Charts!$AB$47:$AB$50", [r[0] for r in lb])
                xml = _replace_numref_values(xml, "Summary_Charts!$AC$47:$AC$50", [r[2] for r in lb])
            elif item.filename.endswith("chart4.xml"):
                lb = metrics["ord_lb_value"]
                xml = _replace_strcache(xml, "Sheet1!$O$29:$O$32", [r[0] for r in lb])
                xml = _replace_numref_values(xml, "Sheet1!$P$29:$P$32", [r[1] for r in lb])
                xml = _replace_numref_values(xml, "Sheet1!$Q$29:$Q$32", [r[2] for r in lb])
            elif item.filename.endswith("chart5.xml"):
                lb = metrics["ord_lb_time"]
                xml = _replace_strcache(xml, "Summary_Charts!$AB$67:$AB$70", [r[0] for r in lb])
                xml = _replace_numref_values(xml, "Summary_Charts!$AC$67:$AC$70", [r[3] for r in lb])
            elif item.filename.endswith("chart6.xml"):
                lb = metrics["q_lb_value"]
                xml = _replace_strcache(xml, "Sheet1!$O$37:$O$40", [r[0] for r in lb])
                xml = _replace_numref_values(xml, "Sheet1!$P$37:$P$40", [r[1] for r in lb])
                xml = _replace_numref_values(xml, "Sheet1!$Q$37:$Q$40", [r[2] for r in lb])
            elif item.filename.endswith("chart7.xml"):
                lb = metrics["q_lb_time"]
                xml = _replace_strcache(xml, "Summary_Charts!$AB$77:$AB$80", [r[0] for r in lb])
                xml = _replace_numref_values(xml, "Summary_Charts!$AC$77:$AC$80", [r[3] for r in lb])
            chart_updates[item.filename] = xml.encode("utf-8")

    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = chart_updates.get(item.filename, zin.read(item.filename))
            zout.writestr(item, data)
    os.replace(tmp, docx_path)


def _get_style(xml: str, cell_ref: str) -> str:
    m = re.search(rf'<c r="{re.escape(cell_ref)}"[^>]*>', xml)
    if not m:
        return ""
    s = re.search(r' s="(\d+)"', m.group(0))
    return s.group(1) if s else ""


def _str_cell(cell_ref: str, text: str, style: str = "") -> str:
    sa = f' s="{style}"' if style else ""
    return f'<c r="{cell_ref}"{sa} t="inlineStr"><is><t>{_xml_escape(text)}</t></is></c>'


def _num_cell(cell_ref: str, value: Any, style: str = "") -> str:
    sa = f' s="{style}"' if style else ""
    return f'<c r="{cell_ref}"{sa}><v>{_num(value)}</v></c>'


def _col_letters(cell_ref: str) -> str:
    return re.match(r"[A-Z]+", cell_ref).group(0)


def _row_num(cell_ref: str) -> int:
    return int(re.search(r"\d+", cell_ref).group(0))


def _patch_cell_in_xml(xml: str, cell_ref: str, cell_xml: str) -> str:
    pattern = re.compile(rf'<c r="{re.escape(cell_ref)}"[^>]*>.*?</c>|<c r="{re.escape(cell_ref)}"[^>]*/>', re.DOTALL)
    if pattern.search(xml):
        return pattern.sub(cell_xml, xml, count=1)
    row = _row_num(cell_ref)
    col = _col_letters(cell_ref)
    def add_cell(m: re.Match) -> str:
        content = m.group(2)
        cells = list(re.finditer(r'<c r="([A-Z]+)\d+"', content))
        pos = next((cm.start() for cm in cells if cm.group(1) > col), len(content))
        return m.group(1) + content[:pos] + cell_xml + content[pos:] + m.group(3)
    return re.sub(rf'(<row r="{row}"[^>]*>)(.*?)(</row>)', add_cell, xml, count=1, flags=re.DOTALL)


def _fix_zip_flags(xlsx_bytes: bytes) -> bytes:
    raw = bytearray(xlsx_bytes)
    for sig, offset in [(b"PK\x03\x04", 6), (b"PK\x01\x02", 8)]:
        pos = 0
        while True:
            pos = bytes(raw).find(sig, pos)
            if pos == -1:
                break
            if struct.unpack_from("<H", raw, pos + offset)[0] == 0:
                struct.pack_into("<H", raw, pos + offset, 6)
            pos += 4
    return bytes(raw)

def _add_percent_style_to_xlsx(xlsx_bytes: bytes) -> Tuple[bytes, str]:
    """Add a 0.00% cell style to an embedded workbook and return (bytes, style_id)."""
    src = zipfile.ZipFile(io.BytesIO(xlsx_bytes), "r")
    styles_path = "xl/styles.xml"
    styles = src.read(styles_path).decode("utf-8")

    fmt_ids = [int(x) for x in re.findall(r'<numFmt numFmtId="(\d+)"', styles)]
    xf_ids = re.findall(r'<xf\b[^>]*/>', re.search(r'<cellXfs[^>]*>.*?</cellXfs>', styles, re.DOTALL).group(0))
    style_id = str(len(xf_ids))
    new_num_fmt_id = str(max(fmt_ids + [163]) + 1)

    if "0.00%" in styles:
        m_fmt = re.search(r'<numFmt numFmtId="(\d+)" formatCode="0\.00%"', styles)
        if m_fmt:
            existing_num_fmt_id = m_fmt.group(1)
            cell_xfs = re.search(r'<cellXfs[^>]*>.*?</cellXfs>', styles, re.DOTALL).group(0)
            xfs = re.findall(r'<xf\b[^>]*/>', cell_xfs)
            for idx, xf in enumerate(xfs):
                if f'numFmtId="{existing_num_fmt_id}"' in xf:
                    src.close()
                    return xlsx_bytes, str(idx)

    numfmt_xml = f'<numFmt numFmtId="{new_num_fmt_id}" formatCode="0.00%"/>'
    if "<numFmts" in styles:
        styles = re.sub(r'<numFmts count="(\d+)">', lambda m: f'<numFmts count="{int(m.group(1)) + 1}">', styles, count=1)
        styles = styles.replace("</numFmts>", numfmt_xml + "</numFmts>", 1)
    else:
        styles = styles.replace("<fonts", f'<numFmts count="1">{numfmt_xml}</numFmts><fonts', 1)

    xf_xml = f'<xf numFmtId="{new_num_fmt_id}" fontId="0" fillId="0" borderId="0" xfId="0" applyNumberFormat="1"/>'
    styles = re.sub(r'<cellXfs count="(\d+)">', lambda m: f'<cellXfs count="{int(m.group(1)) + 1}">', styles, count=1)
    styles = styles.replace("</cellXfs>", xf_xml + "</cellXfs>", 1)

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w") as zout:
        for item in src.infolist():
            data = styles.encode("utf-8") if item.filename == styles_path else src.read(item.filename)
            zout.writestr(item, data)
    src.close()
    return _fix_zip_flags(out.getvalue()), style_id


def _rewrite_xlsx(xlsx_bytes: bytes, sheet_patches: Dict[str, List[Tuple[str, str]]]) -> bytes:
    src = zipfile.ZipFile(io.BytesIO(xlsx_bytes), "r")
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w") as zout:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename in sheet_patches:
                xml = data.decode("utf-8")
                for ref, cell_xml in sheet_patches[item.filename]:
                    xml = _patch_cell_in_xml(xml, ref, cell_xml)
                data = xml.encode("utf-8")
            zout.writestr(item, data)
    src.close()
    return _fix_zip_flags(out.getvalue())


def _make_patches_for_sheet(xlsx_bytes: bytes, sheet_name: str, raw_items: List[Tuple[str, str, Any]]) -> Dict[str, List[Tuple[str, str]]]:
    with zipfile.ZipFile(io.BytesIO(xlsx_bytes), "r") as z:
        xml = z.read(sheet_name).decode("utf-8")
    patches = []
    for cell_ref, kind, value in raw_items:
        style = _get_style(xml, cell_ref)
        if kind == "s":
            patches.append((cell_ref, _str_cell(cell_ref, str(value), style)))
        else:
            patches.append((cell_ref, _num_cell(cell_ref, value, style)))
    return {sheet_name: patches}


def _patch_embedded_workbooks(docx_path: str, metrics: Dict[str, Any]) -> None:
    tmp = docx_path + ".emb.tmp"
    ord_val = metrics["ord_val"]
    q_val = metrics["q_val"]
    total_val = metrics["total_val"] or 1.0
    ord_hrs = metrics["ord_hrs"]
    q_hrs = metrics["q_hrs"]
    all_hrs = metrics["all_hrs"] or 1.0

    workbook_updates: Dict[str, bytes] = {}
    with zipfile.ZipFile(docx_path, "r") as zin:
        for i in range(7):
            name = f'word/embeddings/Microsoft_Excel_Worksheet{("" if i == 0 else i)}.xlsx'
            if name not in zin.namelist():
                continue
            xbytes = zin.read(name)
            if i == 0:  # chart1
                # Add/assign a true percentage number style for the visible data table rows.
                xbytes, pct_style = _add_percent_style_to_xlsx(xbytes)
                with zipfile.ZipFile(io.BytesIO(xbytes), "r") as zx:
                    sheet_xml = zx.read("xl/worksheets/sheet5.xml").decode("utf-8")
                raw_values = {
                    "J3": ord_val, "K3": metrics["active_o_count"], "L3": ord_val / total_val, "M3": ord_hrs, "N3": ord_hrs / all_hrs,
                    "J4": q_val, "K4": metrics["active_q_count"], "L4": q_val / total_val, "M4": q_hrs, "N4": q_hrs / all_hrs,
                    "J5": total_val, "K5": metrics["active_o_count"] + metrics["active_q_count"], "L5": 1, "M5": all_hrs, "N5": 1,
                }
                patches_list = []
                for cell_ref, val in raw_values.items():
                    style = pct_style if cell_ref[0] in {"L", "N"} else _get_style(sheet_xml, cell_ref)
                    patches_list.append((cell_ref, _num_cell(cell_ref, val, style)))
                workbook_updates[name] = _rewrite_xlsx(xbytes, {"xl/worksheets/sheet5.xml": patches_list})
            elif i == 1:  # chart2 frequencies
                items = []
                for idx, val in enumerate(metrics["ord_freq"], start=2):
                    items.append((f"Q{idx}", "n", val))
                for idx, val in enumerate(metrics["q_freq"], start=2):
                    items.append((f"R{idx}", "n", val))
                patches = _make_patches_for_sheet(xbytes, "xl/worksheets/sheet5.xml", items)
                workbook_updates[name] = _rewrite_xlsx(xbytes, patches)
            elif i == 2:  # chart3 order count: sheet6 summary names/counts
                lb = metrics["ord_lb_count"]
                items = []
                for n, row in enumerate(lb):
                    items.append((f"AB{47+n}", "s", row[0]))
                    items.append((f"AC{47+n}", "n", row[2]))
                patches = _make_patches_for_sheet(xbytes, "xl/worksheets/sheet6.xml", items)
                workbook_updates[name] = _rewrite_xlsx(xbytes, patches)
            elif i == 3:  # chart4 order value: sheet5 O29:Q32
                lb = metrics["ord_lb_value"]
                items = []
                for n, row in enumerate(lb):
                    r = 29 + n
                    items += [(f"O{r}", "s", row[0]), (f"P{r}", "n", row[1]), (f"Q{r}", "n", row[2])]
                patches = _make_patches_for_sheet(xbytes, "xl/worksheets/sheet5.xml", items)
                workbook_updates[name] = _rewrite_xlsx(xbytes, patches)
            elif i == 4:  # chart5 order RM time: sheet6 AB67:AC70
                lb = metrics["ord_lb_time"]
                items = []
                for n, row in enumerate(lb):
                    items.append((f"AB{67+n}", "s", row[0]))
                    items.append((f"AC{67+n}", "n", row[3]))
                patches = _make_patches_for_sheet(xbytes, "xl/worksheets/sheet6.xml", items)
                workbook_updates[name] = _rewrite_xlsx(xbytes, patches)
            elif i == 5:  # chart6 quote value: sheet5 O37:Q40
                lb = metrics["q_lb_value"]
                items = []
                for n, row in enumerate(lb):
                    r = 37 + n
                    items += [(f"O{r}", "s", row[0]), (f"P{r}", "n", row[1]), (f"Q{r}", "n", row[2])]
                patches = _make_patches_for_sheet(xbytes, "xl/worksheets/sheet5.xml", items)
                workbook_updates[name] = _rewrite_xlsx(xbytes, patches)
            elif i == 6:  # chart7 quote RM time: sheet6 AB77:AC80
                lb = metrics["q_lb_time"]
                items = []
                for n, row in enumerate(lb):
                    items.append((f"AB{77+n}", "s", row[0]))
                    items.append((f"AC{77+n}", "n", row[3]))
                patches = _make_patches_for_sheet(xbytes, "xl/worksheets/sheet6.xml", items)
                workbook_updates[name] = _rewrite_xlsx(xbytes, patches)

    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = workbook_updates.get(item.filename, zin.read(item.filename))
            compress_type = zipfile.ZIP_STORED if (item.filename.startswith("word/embeddings/") and item.filename.endswith(".xlsx")) else item.compress_type
            zout.writestr(item.filename, data, compress_type=compress_type)
    os.replace(tmp, docx_path)




def _strip_purple_bar_picts(xml: str) -> str:
    """
    Remove the template's old page-body floating purple bars without touching
    charts, logos, or other artwork.

    Word may store these bars in two different ways after python-docx saves the
    file:
      1. Legacy VML: <w:pict><v:rect ... fillcolor="#6f55d7" .../>
      2. DrawingML/WPS: <w:drawing><wp:anchor ...><wp:docPr name="Rectangle 3"/>...

    The new repeating sidebar is then installed in the header layer only.
    """

    def is_legacy_vml_bar(block: str) -> bool:
        b = block.lower()
        return (
            "#6f55d7" in b
            and "width:18pt" in b
            and ("height:842pt" in b or "height:843.4pt" in b)
        )

    def is_drawingml_bar(block: str) -> bool:
        b = block.lower()
        return (
            '<wp:docpr' in b
            and 'name="rectangle 3"' in b
            and 'cx="228600"' in b
            and ('cy="10711180"' in b or 'cy="10693400"' in b or 'cy="10692130"' in b)
            and 'val="accent1"' in b
            and 'val="50000"' in b
        )

    def strip_vml(m: re.Match) -> str:
        block = m.group(0)
        return "" if is_legacy_vml_bar(block) else block

    def strip_drawing(m: re.Match) -> str:
        block = m.group(0)
        return "" if is_drawingml_bar(block) else block

    xml = re.sub(r"<w:pict\b[^>]*>.*?</w:pict>", strip_vml, xml, flags=re.DOTALL)
    xml = re.sub(r"<w:drawing\b[^>]*>.*?</w:drawing>", strip_drawing, xml, flags=re.DOTALL)
    return xml


def _header_bar_paragraph() -> str:
    """A page-relative purple bar placed in the header so Word repeats it on every page."""
    return (
        '<w:p w14:paraId="8B4F5A21" w14:textId="77777777" w:rsidR="00A5655E" w:rsidRDefault="00A5655E">'
        '<w:pPr><w:pStyle w:val="Header"/></w:pPr>'
        '<w:r><w:rPr><w:noProof/></w:rPr>'
        '<w:pict w14:anchorId="DWRLEFTBAR">'
        '<v:rect id="_x0000_s4097" '
        'style="position:absolute;margin-left:0;margin-top:0;width:18pt;height:842pt;z-index:251660288;visibility:visible;'
        'mso-wrap-style:square;mso-width-percent:0;mso-height-percent:0;'
        'mso-wrap-distance-left:0;mso-wrap-distance-top:0;mso-wrap-distance-right:0;mso-wrap-distance-bottom:0;'
        'mso-position-horizontal:absolute;mso-position-horizontal-relative:page;'
        'mso-position-vertical:absolute;mso-position-vertical-relative:page;'
        'mso-width-percent:0;mso-height-percent:0;v-text-anchor:top" '
        'fillcolor="#6f55d7" stroked="f">'
        '<w10:wrap anchorx="page" anchory="page"/><w10:anchorlock/>'
        '</v:rect></w:pict></w:r></w:p>'
    )


def _ensure_header_bar(xml: str) -> str:
    xml = _strip_purple_bar_picts(xml)
    bar = _header_bar_paragraph()
    if "</w:hdr>" not in xml:
        return xml
    return re.sub(r"(<w:hdr\b[^>]*>)", r"\1" + bar, xml, count=1)


def _next_relationship_id(rels_xml: str) -> str:
    ids = [int(x) for x in re.findall(r'Id="rId(\d+)"', rels_xml)]
    return f"rId{max(ids + [0]) + 1}"


def _install_repeating_left_bar(docx_path: str) -> None:
    """
    Replace old per-page floating body bars with one repeating header bar.

    The visual bar definition is intentionally kept the same as the confirmed
    working version: VML rectangle, #6f55d7, 18pt wide, page-relative.
    """
    tmp = docx_path + ".leftbar.tmp"
    with zipfile.ZipFile(docx_path, "r") as zin:
        names = set(zin.namelist())
        doc_xml = zin.read("word/document.xml").decode("utf-8")
        rels_xml = zin.read("word/_rels/document.xml.rels").decode("utf-8")
        content_types = zin.read("[Content_Types].xml").decode("utf-8")
        header1_xml = zin.read("word/header1.xml").decode("utf-8") if "word/header1.xml" in names else None

        doc_xml = _strip_purple_bar_picts(doc_xml)

        if header1_xml is None:
            raise ValueError("Step 2 template must contain word/header1.xml to install the repeating left bar.")
        header1_xml = _ensure_header_bar(header1_xml)

        # Avoid a separate first-page header mode. It can change first-page
        # pagination and push the first pie chart to page 2. The default header
        # bar now repeats on every page.
        doc_xml = doc_xml.replace('<w:titlePg/>', '')
        rels_xml = re.sub(
            r'<Relationship\s+Id="rId\d+"\s+Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/header"\s+Target="header2.xml"\s*/>',
            '',
            rels_xml,
        )
        content_types = content_types.replace(
            '<Override PartName="/word/header2.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"/>',
            '',
        )

        with zipfile.ZipFile(tmp, "w") as zout:
            for item in zin.infolist():
                if item.filename == "word/header2.xml":
                    continue
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = doc_xml.encode("utf-8")
                elif item.filename == "word/_rels/document.xml.rels":
                    data = rels_xml.encode("utf-8")
                elif item.filename == "[Content_Types].xml":
                    data = content_types.encode("utf-8")
                elif item.filename == "word/header1.xml":
                    data = header1_xml.encode("utf-8")
                zout.writestr(item, data)
    os.replace(tmp, docx_path)


def _fit_first_order_count_chart(docx_path: str) -> None:
    """
    Keep the first order-count pie chart on page 1 while restoring its template
    size. The pie chart itself stays at the original Step 2 template extent
    (5384800 x 2263140 EMU). Only the first overview chart above it is slightly
    compacted, which is what allows the pie chart to fit without shrinking.
    """
    tmp = docx_path + ".fitpie.tmp"
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                xml = data.decode("utf-8")

                def patch_chart(m: re.Match) -> str:
                    block = m.group(0)
                    if 'r:id="rId12"' in block:
                        # Slightly reduce only the first overview chart height.
                        return re.sub(
                            r'<wp:extent cx="\d+" cy="\d+"/>',
                            '<wp:extent cx="4572000" cy="2480000"/>',
                            block,
                            count=1,
                        )
                    if 'r:id="rId14"' in block:
                        # Restore original template size for first pie chart.
                        return re.sub(
                            r'<wp:extent cx="\d+" cy="\d+"/>',
                            '<wp:extent cx="5384800" cy="2263140"/>',
                            block,
                            count=1,
                        )
                    return block

                xml = re.sub(r'<w:drawing\b[^>]*>.*?</w:drawing>', patch_chart, xml, flags=re.DOTALL)
                data = xml.encode("utf-8")
            zout.writestr(item, data)
    os.replace(tmp, docx_path)

def _patch_xml_text_and_props(docx_path: str, metrics: Dict[str, Any]) -> None:
    report_date = metrics["date"]
    date_long = report_date.strftime("%d %B %Y").lstrip("0")
    date_upper = date_long.upper()
    date_short = report_date.strftime("%d/%m/%y")
    iso_date = report_date.strftime("%Y-%m-%dT00:00:00")
    dwr = metrics["dwr_no"]
    replacements = {
        "DWR# 2026-90": f"DWR# {dwr}",
        "21/05/26": date_short,
        "21 May 2026": date_long,
        "[PUBLISH DATE]": date_upper,
        "[Publish Date]": date_upper,
        "……": date_long,
        ".....": date_long,
    }
    tmp = docx_path + ".xml.tmp"
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith(".xml") and (item.filename.startswith("word/") or item.filename.startswith("docProps/") or item.filename.startswith("customXml/")):
                try:
                    xml = data.decode("utf-8")
                    for old, new in replacements.items():
                        xml = xml.replace(old, new)
                    if item.filename == "docProps/core.xml":
                        xml = re.sub(r"<cp:contentStatus>.*?</cp:contentStatus>", f"<cp:contentStatus>{dwr}</cp:contentStatus>", xml)
                    if item.filename == "customXml/item1.xml":
                        xml = re.sub(r"<PublishDate>.*?</PublishDate>", f"<PublishDate>{iso_date}</PublishDate>", xml)
                    xml = re.sub(r'<w:date w:fullDate="[^"]+"', f'<w:date w:fullDate="{iso_date}Z"', xml)
                    data = xml.encode("utf-8")
                except UnicodeDecodeError:
                    pass
            zout.writestr(item, data)
    os.replace(tmp, docx_path)


def generate_dwr_step2(excel_path: str, template_path: str | None = None, output_dir: str | None = None) -> str:
    """Generate the Step 2 DWR DOCX and return the output path."""
    template_path = template_path or DEFAULT_TEMPLATE
    output_dir = output_dir or DEFAULT_OUTPUT_DIR
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Step 2 template not found: {template_path}")
    os.makedirs(output_dir, exist_ok=True)

    data = _read_excel_data(excel_path)
    metrics = _compute_metrics(data)

    dwr_no = metrics["dwr_no"]
    d = metrics["date"]
    out_name = f"DWR_{dwr_no}_{d.strftime('%d-%b-%Y')}-Step_2.docx"
    output_path = os.path.join(output_dir, out_name)

    shutil.copyfile(template_path, output_path)
    _patch_analysis_tables(output_path, metrics)
    _patch_xml_text_and_props(output_path, metrics)
    _patch_charts(output_path, metrics)
    _patch_embedded_workbooks(output_path, metrics)
    _install_repeating_left_bar(output_path)
    _fit_first_order_count_chart(output_path)
    return output_path


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Generate SDL DWR Step 2 report DOCX.")
    parser.add_argument("excel_path")
    parser.add_argument("--template", default=None)
    parser.add_argument("--output-dir", default=None)
    args = parser.parse_args()
    print(generate_dwr_step2(args.excel_path, template_path=args.template, output_dir=args.output_dir))
