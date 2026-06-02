# -*- coding: utf-8 -*-
"""
sdl_dwr_step3_generator.py

Generate Step 3 SDL Daily Work Report DOCX from the daily Excel file.
Designed for a small Flask app on Namecheap/cPanel shared hosting.

Required packages:
    openpyxl
    python-docx

Expected server files in the same folder as this script/app:
    Daily Work Report Step 3 Template - enTop v1.0.0.docx

Public function:
    generate_dwr_step3(excel_path, template_path=None, output_dir=None) -> str
"""

from __future__ import annotations

import html
import io
import math
import os
import re
import shutil
import statistics
import struct
import zipfile
from collections import defaultdict
from datetime import date, datetime
from typing import Any, Dict, Iterable, List, Tuple

from docx import Document
from docx.table import _Cell
from openpyxl import load_workbook

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_TEMPLATE = os.path.join(BASE_DIR, "Daily Work Report Step 3 Template - enTop v1.0.0.docx")
DEFAULT_OUTPUT_DIR = os.path.join(BASE_DIR, "generated_reports")

ONE_HOUR_BINS = [1, 2, 3, 4, 5, 6, 7, 8]
HALF_HOUR_BINS = [0.5, 1.0, 1.5, 2.0, 2.5, 3.0, 3.5, 4.0, 4.5, 5.0, 5.5, 6.0, 6.5, 7.0]


def _safe_float(value: Any, default: float = 0.0) -> float:
    if value is None or value == "":
        return default
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _normalize_job_no(raw: Any) -> str:
    if raw is None:
        return ""
    if isinstance(raw, (int, float)):
        return str(int(raw))
    return str(raw).strip()


def _normalize_customer(raw: Any) -> str:
    return " ".join(str(raw).split()) if raw else ""


def _normalize_status(raw: Any) -> str:
    return str(raw).strip().upper() if raw is not None else ""


def _normalize_grp(raw: Any) -> str:
    return str(raw).strip() if raw is not None else ""


def _excel_serial_or_date(value: Any) -> date:
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    try:
        return datetime.strptime(str(value)[:10], "%Y-%m-%d").date()
    except Exception:
        return date.today()


def _fmt_k(v: float) -> str:
    return f"${v / 1000:.1f}k"


def _fmt_hrs(v: float) -> str:
    return f"{v:.2f}"


def _read_excel_data(excel_path: str) -> Dict[str, Any]:
    wb = load_workbook(excel_path, data_only=True)
    if "Summary" not in wb.sheetnames or "Production_Center" not in wb.sheetnames:
        raise ValueError("Input workbook must contain 'Summary' and 'Production_Center' sheets.")

    ws_sum = wb["Summary"]
    ws_pc = wb["Production_Center"]
    dwr_no = str(ws_sum["B1"].value).strip()
    report_date = _excel_serial_or_date(ws_sum["C1"].value)

    jobs: List[Dict[str, Any]] = []
    for r in range(8, min(ws_pc.max_row, 108) + 1):
        job_type = str(ws_pc.cell(r, 4).value or "").strip().upper()
        job_no = _normalize_job_no(ws_pc.cell(r, 5).value)
        if not job_type and not job_no:
            continue
        grp_work = _normalize_grp(ws_pc.cell(r, 15).value)
        if grp_work.lower() == "continued":
            continue
        jobs.append({
            "row": r,
            "type": job_type,
            "job_no": job_no,
            "value": _safe_float(ws_pc.cell(r, 11).value),
            "customer": _normalize_customer(ws_pc.cell(r, 12).value),
            "grp_work": grp_work,
            "rm_time": _safe_float(ws_pc.cell(r, 21).value),
            "status": _normalize_status(ws_pc.cell(r, 25).value),
        })

    active = [j for j in jobs if j["status"] in {"C", "IP"}]
    active_o = [j for j in active if j["type"] == "O"]
    active_q = [j for j in active if j["type"] == "Q"]

    return {
        "dwr_no": dwr_no,
        "date": report_date,
        "jobs": jobs,
        "active": active,
        "active_o": active_o,
        "active_q": active_q,
        "ord_val": sum(j["value"] for j in active_o),
        "q_val": sum(j["value"] for j in active_q),
        "ord_hrs": sum(j["rm_time"] for j in active_o),
        "q_hrs": sum(j["rm_time"] for j in active_q),
    }


def _make_leaderboard(jobs: List[Dict[str, Any]], sort_key: str, top_n: int = 3) -> List[Tuple[str, float, int, float]]:
    cust = defaultdict(lambda: {"value": 0.0, "count": 0, "time": 0.0})
    for j in jobs:
        customer = j["customer"] or "UNKNOWN CUSTOMER"
        cust[customer]["value"] += j["value"]
        cust[customer]["count"] += 1
        cust[customer]["time"] += j["rm_time"]
    ranked = sorted(cust.items(), key=lambda x: (-x[1][sort_key], x[0]))
    top = ranked[:top_n]
    rest = ranked[top_n:]
    rows = [(c, d["value"], int(d["count"]), d["time"]) for c, d in top]
    while len(rows) < top_n:
        rows.append(("", 0.0, 0, 0.0))
    rows.append((
        "OTHER CUSTOMERS",
        sum(d["value"] for _, d in rest),
        int(sum(d["count"] for _, d in rest)),
        sum(d["time"] for _, d in rest),
    ))
    return rows


def _bin_counts(values: Iterable[float], uppers: List[float]) -> List[int]:
    vals = [v for v in values if v is not None and v > 0]
    counts = []
    prev = 0.0
    for high in uppers:
        counts.append(sum(1 for v in vals if prev < v <= high))
        prev = high
    return counts


def _normal_cdf(z: float) -> float:
    return 0.5 * (1.0 + math.erf(z / math.sqrt(2.0)))


def _lognormal_expected(values: Iterable[float], uppers: List[float]) -> List[float]:
    vals = [v for v in values if v and v > 0]
    if not vals:
        return [0.0 for _ in uppers]
    logs = [math.log(v) for v in vals]
    mu = statistics.mean(logs)
    sigma = statistics.stdev(logs) if len(logs) > 1 else 0.000001
    if sigma <= 0:
        sigma = 0.000001
    n = len(vals)
    expected = []
    prev = 0.0
    for high in uppers:
        low_cdf = 0.0 if prev <= 0 else _normal_cdf((math.log(prev) - mu) / sigma)
        high_cdf = _normal_cdf((math.log(high) - mu) / sigma)
        expected.append(max(0.0, n * (high_cdf - low_cdf)))
        prev = high
    return expected


def _desc_stats(times: List[float]) -> Dict[str, str]:
    s = sorted([t for t in times if t is not None and t > 0])
    if not s:
        return {
            "count": "0", "sum": "0.00 hrs", "mean": "0.00 hrs", "median": "0.00 hrs",
            "min": "0.00 hrs", "max": "0.00 hrs", "range": "0.00 hrs",
            "q1": "0.00 hrs", "q3": "0.00 hrs", "stddev": "0.00 hrs",
        }
    n = len(s)
    q1 = statistics.quantiles(s, n=4)[0] if n >= 2 else s[0]
    q3 = statistics.quantiles(s, n=4)[2] if n >= 2 else s[0]
    return {
        "count": str(n),
        "sum": f"{sum(s):.2f} hrs",
        "mean": f"{statistics.mean(s):.2f} hrs",
        "median": f"{statistics.median(s):.2f} hrs",
        "min": f"{min(s):.2f} hrs",
        "max": f"{max(s):.2f} hrs",
        "range": f"{max(s) - min(s):.2f} hrs",
        "q1": f"{q1:.2f} hrs",
        "q3": f"{q3:.2f} hrs",
        "stddev": f"{statistics.stdev(s):.2f} hrs" if n > 1 else "0.00 hrs",
    }


def _compute_metrics(data: Dict[str, Any]) -> Dict[str, Any]:
    active = data["active"]
    active_o = data["active_o"]
    active_q = data["active_q"]
    ord_val = data["ord_val"]
    q_val = data["q_val"]
    ord_hrs = data["ord_hrs"]
    q_hrs = data["q_hrs"]
    all_times = [j["rm_time"] for j in active if j["rm_time"] > 0]
    o_times = [j["rm_time"] for j in active_o if j["rm_time"] > 0]
    q_times = [j["rm_time"] for j in active_q if j["rm_time"] > 0]

    metrics = dict(data)
    metrics.update({
        "ord_lb_count": _make_leaderboard(active_o, "count"),
        "ord_lb_value": _make_leaderboard(active_o, "value"),
        "ord_lb_time": _make_leaderboard(active_o, "time"),
        "q_lb_count": _make_leaderboard(active_q, "count"),
        "q_lb_value": _make_leaderboard(active_q, "value"),
        "q_lb_time": _make_leaderboard(active_q, "time"),
        "stats_all": _desc_stats(all_times),
        "stats_o": _desc_stats(o_times),
        "stats_q": _desc_stats(q_times),
        "hist_1hr": _bin_counts(all_times, ONE_HOUR_BINS),
        "hist_half": _bin_counts(all_times, HALF_HOUR_BINS),
        "lognorm_expected": _lognormal_expected(all_times, HALF_HOUR_BINS),
        "ord_val": ord_val,
        "q_val": q_val,
        "ord_hrs": ord_hrs,
        "q_hrs": q_hrs,
    })
    return metrics


def _set_cell_text(cell: _Cell, text: Any) -> None:
    text = "" if text is None else str(text)
    if not cell.paragraphs:
        cell.text = text
        return
    p = cell.paragraphs[0]
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for run in p.runs[1:]:
        run.text = ""
    for extra_p in cell.paragraphs[1:]:
        for run in extra_p.runs:
            run.text = ""


def _replace_in_paragraphs(paragraphs: Iterable[Any], replacements: Dict[str, str]) -> None:
    for p in paragraphs:
        full = "".join(run.text for run in p.runs)
        new_full = full
        for old, new in replacements.items():
            new_full = new_full.replace(old, new)
        if new_full != full and p.runs:
            p.runs[0].text = new_full
            for r in p.runs[1:]:
                r.text = ""


def _replace_everywhere(doc: Document, replacements: Dict[str, str]) -> None:
    _replace_in_paragraphs(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs, replacements)
    for section in doc.sections:
        for story in (section.header, section.footer):
            _replace_in_paragraphs(story.paragraphs, replacements)
            for table in story.tables:
                for row in table.rows:
                    for cell in row.cells:
                        _replace_in_paragraphs(cell.paragraphs, replacements)


def _patch_doc_tables(docx_path: str, metrics: Dict[str, Any]) -> None:
    doc = Document(docx_path)
    d = metrics["date"]
    date_long = d.strftime("%d %B %Y").lstrip("0")
    date_short = d.strftime("%d/%m/%y")
    dwr = metrics["dwr_no"]
    replacements = {
        "DWR# 2026-90": f"DWR# {dwr}",
        "21/05/26": date_short,
        "21 May 2026": date_long,
    }
    _replace_everywhere(doc, replacements)

    ord_lb_count = metrics["ord_lb_count"]
    ord_lb_value = metrics["ord_lb_value"]
    ord_lb_time = metrics["ord_lb_time"]
    q_lb_count = metrics["q_lb_count"]
    q_lb_value = metrics["q_lb_value"]
    q_lb_time = metrics["q_lb_time"]

    active_o_count = len(metrics["active_o"])
    active_q_count = len(metrics["active_q"])
    ord_val = metrics["ord_val"] or 1.0
    q_val = metrics["q_val"] or 1.0
    ord_hrs = metrics["ord_hrs"] or 1.0
    q_hrs = metrics["q_hrs"] or 1.0

    def mini_rows(lb_count, lb_value, lb_time, total_count, total_value, total_time, kind):
        rows = []
        for i in range(3):
            rows.append([
                lb_count[i][0], str(int(lb_count[i][2])),
                lb_value[i][0], _fmt_k(lb_value[i][1]),
                lb_time[i][0], _fmt_hrs(lb_time[i][3]),
            ])
        cnt_total = sum(r[2] for r in lb_count[:3])
        val_total = sum(r[1] for r in lb_value[:3])
        time_total = sum(r[3] for r in lb_time[:3])
        cnt_pct = int((cnt_total / total_count * 100) if total_count else 0)
        val_pct = int((val_total / total_value * 100) if total_value else 0)
        time_pct = int((time_total / total_time * 100) if total_time else 0)
        rows.append([
            f"Totals ({cnt_pct}% {kind})", str(int(cnt_total)) if cnt_total else "-",
            f"Totals ({val_pct}% Value)", _fmt_k(val_total),
            f"Totals ({time_pct}% Hrs)", _fmt_hrs(time_total),
        ])
        return rows

    # Template tables: 1 = orders mini-table, 3 = quotes mini-table, 4 = descriptive stats.
    if len(doc.tables) >= 5:
        o_rows = mini_rows(ord_lb_count, ord_lb_value, ord_lb_time, active_o_count, ord_val, ord_hrs, "orders")
        t = doc.tables[1]
        for r_idx, vals in enumerate(o_rows, start=1):
            for c_idx, val in enumerate(vals):
                _set_cell_text(t.cell(r_idx, c_idx), val)

        q_rows = mini_rows(q_lb_count, q_lb_value, q_lb_time, active_q_count, q_val, q_hrs, "quotes")
        t = doc.tables[3]
        for r_idx, vals in enumerate(q_rows, start=1):
            for c_idx, val in enumerate(vals):
                _set_cell_text(t.cell(r_idx, c_idx), val)

        stat_keys = ["count", "sum", "mean", "median", "min", "max", "range", "q1", "q3", "stddev"]
        t = doc.tables[4]
        for r_idx, key in enumerate(stat_keys, start=1):
            _set_cell_text(t.cell(r_idx, 1), metrics["stats_all"][key])
            _set_cell_text(t.cell(r_idx, 2), metrics["stats_o"][key])
            _set_cell_text(t.cell(r_idx, 3), metrics["stats_q"][key])

    doc.save(docx_path)


def _xml_escape(v: Any) -> str:
    return html.escape(str(v), quote=False)


def _num(v: Any) -> str:
    f = _safe_float(v)
    return f"{f:.15g}"


def _replace_numref_values(xml: str, formula: str, new_vals: List[Any]) -> str:
    def replace_block(m: re.Match) -> str:
        block = m.group(0)
        if f"<c:f>{formula}</c:f>" not in block:
            return block
        fc = re.search(r"<c:formatCode>[^<]*</c:formatCode>", block)
        fmt = fc.group(0) if fc else ""
        pts_xml = f'<c:ptCount val="{len(new_vals)}"/>' + "".join(
            f'<c:pt idx="{i}"><c:v>{_num(v)}</c:v></c:pt>' for i, v in enumerate(new_vals)
        )
        return re.sub(r"<c:numCache>.*?</c:numCache>", f"<c:numCache>{fmt}{pts_xml}</c:numCache>", block, flags=re.DOTALL)
    return re.sub(r"<c:numRef>.*?</c:numRef>", replace_block, xml, flags=re.DOTALL)


def _patch_charts(docx_path: str, metrics: Dict[str, Any]) -> None:
    chart_updates: Dict[str, bytes] = {}
    tmp = docx_path + ".charts.tmp"
    with zipfile.ZipFile(docx_path, "r") as zin:
        for item in zin.infolist():
            if item.filename == "word/charts/chart1.xml":
                xml = zin.read(item.filename).decode("utf-8")
                xml = _replace_numref_values(xml, "Sheet1!$B$2:$B$9", metrics["hist_1hr"])
                chart_updates[item.filename] = xml.encode("utf-8")
            elif item.filename == "word/charts/chart2.xml":
                xml = zin.read(item.filename).decode("utf-8")
                xml = _replace_numref_values(xml, "Descriptive_Stats!$Z$15:$Z$28", metrics["hist_half"])
                xml = _replace_numref_values(xml, "Descriptive_Stats!$AA$15:$AA$28", [round(v, 4) for v in metrics["lognorm_expected"]])
                chart_updates[item.filename] = xml.encode("utf-8")
            elif item.filename == "word/charts/chartEx1.xml":
                xml = zin.read(item.filename).decode("utf-8")
                q_sorted = sorted([j["rm_time"] for j in metrics["active_q"] if j["rm_time"] > 0])
                o_sorted = sorted([j["rm_time"] for j in metrics["active_o"] if j["rm_time"] > 0])
                labels = ["Q"] * len(q_sorted) + ["O"] * len(o_sorted)
                values = q_sorted + o_sorted
                total_pts = len(labels)
                last_row = 1 + total_pts
                str_pts = "".join(f'<cx:pt idx="{i}">{_xml_escape(lbl)}</cx:pt>' for i, lbl in enumerate(labels))
                num_pts = "".join(f'<cx:pt idx="{i}">{_num(val)}</cx:pt>' for i, val in enumerate(values))
                new_str_lvl = f'<cx:lvl ptCount="{total_pts}">{str_pts}</cx:lvl>'
                new_num_lvl = f'<cx:lvl ptCount="{total_pts}" formatCode="General">{num_pts}</cx:lvl>'
                xml = re.sub(r'(<cx:strDim[^>]*>.*?<cx:f>[^<]*</cx:f>)<cx:lvl[^>]*>.*?</cx:lvl>', lambda m: m.group(1) + new_str_lvl, xml, flags=re.DOTALL)
                xml = re.sub(r'(<cx:numDim[^>]*>.*?<cx:f>[^<]*</cx:f>)<cx:lvl[^>]*>.*?</cx:lvl>', lambda m: m.group(1) + new_num_lvl, xml, flags=re.DOTALL)
                xml = re.sub(r'<cx:f>Sheet1!\$A\$2:\$A\$\d+</cx:f>', f'<cx:f>Sheet1!$A$2:$A${last_row}</cx:f>', xml)
                xml = re.sub(r'<cx:f>Sheet1!\$B\$2:\$B\$\d+</cx:f>', f'<cx:f>Sheet1!$B$2:$B${last_row}</cx:f>', xml)
                chart_updates[item.filename] = xml.encode("utf-8")

    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = chart_updates.get(item.filename, zin.read(item.filename))
            zout.writestr(item, data)
    os.replace(tmp, docx_path)


def _col_letters(cell_ref: str) -> str:
    return re.match(r"[A-Z]+", cell_ref).group(0)


def _row_num(cell_ref: str) -> int:
    return int(re.search(r"\d+", cell_ref).group(0))


def _patch_cell_in_xml(xml: str, cell_ref: str, cell_xml: str) -> str:
    pattern = re.compile(rf'<c r="{re.escape(cell_ref)}"[^>]*>.*?</c>|<c r="{re.escape(cell_ref)}"[^>]*/>', re.DOTALL)
    if pattern.search(xml):
        return pattern.sub(cell_xml, xml, count=1)
    row = _row_num(cell_ref)
    col = _col_letters(cell_ref)
    def add_cell(m: re.Match) -> str:
        content = m.group(2)
        cells = list(re.finditer(r'<c r="([A-Z]+)\d+"', content))
        pos = next((cm.start() for cm in cells if cm.group(1) > col), len(content))
        return m.group(1) + content[:pos] + cell_xml + content[pos:] + m.group(3)
    return re.sub(rf'(<row r="{row}"[^>]*>)(.*?)(</row>)', add_cell, xml, count=1, flags=re.DOTALL)


def _num_cell(cell_ref: str, value: Any) -> str:
    return f'<c r="{cell_ref}"><v>{_num(value)}</v></c>'


def _fix_zip_flags(xlsx_bytes: bytes) -> bytes:
    raw = bytearray(xlsx_bytes)
    for sig, offset in [(b"PK\x03\x04", 6), (b"PK\x01\x02", 8)]:
        pos = 0
        while True:
            pos = bytes(raw).find(sig, pos)
            if pos == -1:
                break
            if struct.unpack_from("<H", raw, pos + offset)[0] == 0:
                struct.pack_into("<H", raw, pos + offset, 6)
            pos += 4
    return bytes(raw)


def _rewrite_xlsx(xlsx_bytes: bytes, sheet_patches: Dict[str, List[Tuple[str, str]]], rebuild_sheetdata: Dict[str, str] | None = None, dimensions: Dict[str, str] | None = None) -> bytes:
    rebuild_sheetdata = rebuild_sheetdata or {}
    dimensions = dimensions or {}
    src = zipfile.ZipFile(io.BytesIO(xlsx_bytes), "r")
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w") as zout:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename in sheet_patches or item.filename in rebuild_sheetdata or item.filename in dimensions:
                xml = data.decode("utf-8")
                if item.filename in dimensions:
                    xml = re.sub(r'<dimension ref="[^"]+"/>', f'<dimension ref="{dimensions[item.filename]}"/>', xml)
                if item.filename in rebuild_sheetdata:
                    xml = re.sub(r'<sheetData>.*?</sheetData>', f'<sheetData>{rebuild_sheetdata[item.filename]}</sheetData>', xml, flags=re.DOTALL)
                for ref, cell_xml in sheet_patches.get(item.filename, []):
                    xml = _patch_cell_in_xml(xml, ref, cell_xml)
                data = xml.encode("utf-8")
            zout.writestr(item, data)
    src.close()
    return _fix_zip_flags(out.getvalue())


def _patch_embedded_workbooks(docx_path: str, metrics: Dict[str, Any]) -> None:
    workbook_updates: Dict[str, bytes] = {}
    tmp = docx_path + ".emb.tmp"
    with zipfile.ZipFile(docx_path, "r") as zin:
        hist_wb = "word/embeddings/Microsoft_Excel_Worksheet.xlsx"
        if hist_wb in zin.namelist():
            xbytes = zin.read(hist_wb)
            patches = [(f"B{idx}", _num_cell(f"B{idx}", val)) for idx, val in enumerate(metrics["hist_1hr"], start=2)]
            workbook_updates[hist_wb] = _rewrite_xlsx(xbytes, {"xl/worksheets/sheet1.xml": patches})

        box_wb = "word/embeddings/Microsoft_Excel_Worksheet1.xlsx"
        if box_wb in zin.namelist():
            xbytes = zin.read(box_wb)
            q_sorted = sorted([j["rm_time"] for j in metrics["active_q"] if j["rm_time"] > 0])
            o_sorted = sorted([j["rm_time"] for j in metrics["active_o"] if j["rm_time"] > 0])
            last_row = 1 + len(q_sorted) + len(o_sorted)
            rows_xml = '<row r="1" spans="1:2" x14ac:dyDescent="0.35"><c r="B1" t="s"><v>0</v></c></row>'
            row_num = 2
            for t in q_sorted:
                rows_xml += f'<row r="{row_num}" spans="1:2" x14ac:dyDescent="0.35"><c r="A{row_num}" t="s"><v>1</v></c><c r="B{row_num}"><v>{_num(t)}</v></c></row>'
                row_num += 1
            for t in o_sorted:
                rows_xml += f'<row r="{row_num}" spans="1:2" x14ac:dyDescent="0.35"><c r="A{row_num}" t="s"><v>2</v></c><c r="B{row_num}"><v>{_num(t)}</v></c></row>'
                row_num += 1
            workbook_updates[box_wb] = _rewrite_xlsx(
                xbytes,
                {},
                rebuild_sheetdata={"xl/worksheets/sheet1.xml": rows_xml},
                dimensions={"xl/worksheets/sheet1.xml": f"A1:B{last_row}"},
            )

    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = workbook_updates.get(item.filename, zin.read(item.filename))
            compress_type = zipfile.ZIP_STORED if (item.filename.startswith("word/embeddings/") and item.filename.endswith(".xlsx")) else item.compress_type
            zout.writestr(item.filename, data, compress_type=compress_type)
    os.replace(tmp, docx_path)


def _patch_xml_text_and_props(docx_path: str, metrics: Dict[str, Any]) -> None:
    d = metrics["date"]
    date_long = d.strftime("%d %B %Y").lstrip("0")
    date_short = d.strftime("%d/%m/%y")
    iso_date = d.strftime("%Y-%m-%dT00:00:00")
    dwr = metrics["dwr_no"]
    replacements = {
        "DWR# 2026-90": f"DWR# {dwr}",
        "21/05/26": date_short,
        "21 May 2026": date_long,
    }
    tmp = docx_path + ".xml.tmp"
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith(".xml") and (item.filename.startswith("word/") or item.filename.startswith("docProps/") or item.filename.startswith("customXml/")):
                try:
                    xml = data.decode("utf-8")
                    for old, new in replacements.items():
                        xml = xml.replace(old, new)
                    if item.filename == "docProps/core.xml":
                        xml = re.sub(r"<cp:contentStatus>.*?</cp:contentStatus>", f"<cp:contentStatus>{dwr}</cp:contentStatus>", xml)
                    if item.filename == "customXml/item1.xml":
                        xml = re.sub(r"<PublishDate>.*?</PublishDate>", f"<PublishDate>{iso_date}</PublishDate>", xml)
                    xml = re.sub(r'<w:date w:fullDate="[^"]+"', f'<w:date w:fullDate="{iso_date}Z"', xml)
                    data = xml.encode("utf-8")
                except UnicodeDecodeError:
                    pass
            zout.writestr(item, data)
    os.replace(tmp, docx_path)


def generate_dwr_step3(excel_path: str, template_path: str | None = None, output_dir: str | None = None) -> str:
    """Generate the Step 3 DWR DOCX and return the output path."""
    template_path = template_path or DEFAULT_TEMPLATE
    output_dir = output_dir or DEFAULT_OUTPUT_DIR
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Step 3 template not found: {template_path}")
    os.makedirs(output_dir, exist_ok=True)

    data = _read_excel_data(excel_path)
    metrics = _compute_metrics(data)

    dwr_no = metrics["dwr_no"]
    d = metrics["date"]
    out_name = f"DWR_{dwr_no}_{d.strftime('%d-%b-%Y')}-Step_3.docx"
    output_path = os.path.join(output_dir, out_name)

    shutil.copyfile(template_path, output_path)
    _patch_doc_tables(output_path, metrics)
    _patch_xml_text_and_props(output_path, metrics)
    _patch_charts(output_path, metrics)
    _patch_embedded_workbooks(output_path, metrics)
    return output_path


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Generate SDL DWR Step 3 report DOCX.")
    parser.add_argument("excel_path")
    parser.add_argument("--template", default=None)
    parser.add_argument("--output-dir", default=None)
    args = parser.parse_args()
    print(generate_dwr_step3(args.excel_path, template_path=args.template, output_dir=args.output_dir))
